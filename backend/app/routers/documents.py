from datetime import datetime
from typing import Annotated

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, Query
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pypdf
import pdfplumber
import docx
import io

from app.config import settings
from app.database import get_db
from app.models.document_summary import DocumentSummary
from app.models.folder_settings import FolderSettings
from app.models.folder_share import FolderShare
from app.models.user import User
from app.services.auth import get_current_user
from app.services.llm import generate_summary
from app.services.vector_store import (
    add_documents,
    delete_source,
    get_document_chunks,
    get_files_in_folder,
    list_sources_by_folder,
    move_source,
    rename_source,
)
from pydantic import BaseModel
from sqlalchemy.orm import Session

router = APIRouter(prefix="/api/documents", tags=["documents"])

ALLOWED_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _parse_pdf(data: bytes) -> str:
    parts = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if text and text.strip():
                    parts.append(text)
                for table in page.extract_tables():
                    for row in table:
                        cells = [str(c).strip() for c in row if c and str(c).strip()]
                        if cells:
                            parts.append("  ".join(cells))
    except Exception:
        reader = pypdf.PdfReader(io.BytesIO(data))
        parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts)


def _parse_docx(data: bytes) -> str:
    doc = docx.Document(io.BytesIO(data))
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    for section in doc.sections:
        for hf in [
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ]:
            try:
                if hf and not hf.is_linked_to_previous:
                    for p in hf.paragraphs:
                        if p.text.strip():
                            parts.append(p.text)
            except Exception:
                pass

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for txbx in doc.element.body.iter(f"{{{WNS}}}txbxContent"):
        for p_elem in txbx.iter(f"{{{WNS}}}p"):
            line = "".join(t.text for t in p_elem.iter(f"{{{WNS}}}t") if t.text).strip()
            if line:
                parts.append(line)

    try:
        footnotes_part = doc.part.footnotes
        if footnotes_part:
            for fn in footnotes_part._element.iter(f"{{{WNS}}}p"):
                line = "".join(t.text for t in fn.iter(f"{{{WNS}}}t") if t.text).strip()
                if line:
                    parts.append(line)
    except Exception:
        pass

    return "\n".join(filter(None, parts))


@router.get("")
async def get_documents(
    current_user: Annotated[User, Depends(get_current_user)],
    db: Annotated[Session, Depends(get_db)],
):
    """Return user's own folders + folders shared with them."""
    own_folders = list_sources_by_folder(current_user.id)

    # Build folder visibility map
    settings_rows = (
        db.query(FolderSettings)
        .filter(FolderSettings.user_id == current_user.id)
        .all()
    )
    folder_visibility: dict[str, bool] = {row.folder_name: row.is_public for row in settings_rows}

    shares = (
        db.query(FolderShare)
        .filter(FolderShare.grantee_id == current_user.id)
        .all()
    )
    shared = []
    for s in shares:
        owner = db.get(User, s.owner_id)
        if owner:
            files = get_files_in_folder(s.owner_id, s.folder_name)
            shared.append({
                "owner": owner.username,
                "owner_id": s.owner_id,
                "folder": s.folder_name,
                "files": files,
            })

    return {"folders": own_folders, "folder_settings": folder_visibility, "shared": shared}


class MoveRequest(BaseModel):
    folder: str


@router.patch("/{filename}/folder")
async def move_document(
    filename: str,
    body: MoveRequest,
    current_user: Annotated[User, Depends(get_current_user)],
):
    updated = move_source(filename, body.folder, current_user.id)
    if updated == 0:
        raise HTTPException(status_code=404, detail="Document not found.")
    return {"filename": filename, "folder": body.folder, "updated_chunks": updated}


class RenameRequest(BaseModel):
    new_name: str


@router.patch("/{filename}/rename")
async def rename_document(
    filename: str,
    body: RenameRequest,
    current_user: Annotated[User, Depends(get_current_user)],
):
    new_name = body.new_name.strip()
    if not new_name:
        raise HTTPException(status_code=400, detail="新檔名不能為空。")
    updated = rename_source(filename, new_name, current_user.id)
    if updated == 0:
        raise HTTPException(status_code=404, detail="Document not found.")
    return {"old_name": filename, "new_name": new_name, "updated_chunks": updated}


@router.delete("/{filename}")
async def delete_document(
    filename: str,
    current_user: Annotated[User, Depends(get_current_user)],
):
    deleted = delete_source(filename, current_user.id)
    if deleted == 0:
        raise HTTPException(status_code=404, detail="Document not found.")
    return {"filename": filename, "deleted_chunks": deleted}


@router.post("/{filename}/summarize")
async def summarize_document(
    filename: str,
    current_user: Annotated[User, Depends(get_current_user)],
    db: Annotated[Session, Depends(get_db)],
    force: bool = Query(False, description="Force regeneration even if a cached summary exists"),
):
    """Generate (or return cached) AI summary for a document."""
    # Return cached summary unless force-regeneration is requested
    if not force:
        cached = (
            db.query(DocumentSummary)
            .filter_by(user_id=current_user.id, filename=filename)
            .first()
        )
        if cached:
            return {"summary": cached.summary, "cached": True, "created_at": cached.created_at}

    chunks = get_document_chunks(filename, current_user.id)
    if not chunks:
        raise HTTPException(status_code=404, detail="Document not found in knowledge base.")

    try:
        summary = generate_summary(filename, chunks)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))

    # Upsert cache
    existing = (
        db.query(DocumentSummary)
        .filter_by(user_id=current_user.id, filename=filename)
        .first()
    )
    if existing:
        existing.summary = summary
        existing.created_at = datetime.utcnow()
    else:
        db.add(DocumentSummary(user_id=current_user.id, filename=filename, summary=summary))
    db.commit()

    return {"summary": summary, "cached": False, "created_at": datetime.utcnow()}


@router.post("/upload")
async def upload_document(
    current_user: Annotated[User, Depends(get_current_user)],
    db: Annotated[Session, Depends(get_db)],
    file: UploadFile = File(...),
    folder: str = Form("未分類"),
    is_public: bool = Form(False),
):
    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    data = await file.read()

    if file.content_type == "application/pdf":
        text = _parse_pdf(data)
    else:
        text = _parse_docx(data)

    if not text.strip():
        raise HTTPException(status_code=422, detail="Could not extract text from file.")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    chunks = splitter.split_text(text)

    documents = [
        Document(
            page_content=chunk,
            metadata={"source": file.filename, "folder": folder, "user_id": current_user.id},
        )
        for chunk in chunks
    ]

    ids = add_documents(documents)

    # Upsert folder visibility setting
    existing_setting = (
        db.query(FolderSettings)
        .filter_by(user_id=current_user.id, folder_name=folder)
        .first()
    )
    if existing_setting:
        existing_setting.is_public = is_public
    else:
        db.add(FolderSettings(user_id=current_user.id, folder_name=folder, is_public=is_public))
    db.commit()

    return {"filename": file.filename, "folder": folder, "chunks": len(ids), "ids": ids, "is_public": is_public}
